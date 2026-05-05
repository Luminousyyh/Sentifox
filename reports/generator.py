"""
报告生成模块
生成 Word 格式的舆情分析报告
"""
import os
import io
import base64
from datetime import datetime
from typing import List, Dict, Any, Optional

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

from config import REPORTS_DIR
from utils.database import (
    get_posts, get_post_count, get_sentiment_distribution,
    get_platform_distribution, get_trend_by_time, get_alerts, get_edges
)
from analysis.trend_analysis import calculate_sentiment_index
from graph_analysis.graph_builder import build_graph_from_posts, get_graph_stats
from graph_analysis.influencer_detection import detect_influencers


def _add_heading(doc: Document, text: str, level: int = 1):
    """添加标题"""
    heading = doc.add_heading(level=level)
    run = heading.add_run(text)
    run.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)
    if level == 1:
        run.font.size = Pt(18)
    elif level == 2:
        run.font.size = Pt(14)
    return heading


def _add_kpi_table(doc: Document, kpis: Dict[str, Any]):
    """添加 KPI 表格"""
    table = doc.add_table(rows=1, cols=2)
    table.style = "Light Grid Accent 1"
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "指标"
    hdr_cells[1].text = "数值"

    for key, val in kpis.items():
        row_cells = table.add_row().cells
        row_cells[0].text = str(key)
        row_cells[1].text = str(val)

    doc.add_paragraph()


def generate_report(
    title: str = "舆情分析报告",
    platforms: Optional[List[str]] = None,
    start_time: Optional[datetime] = None,
    end_time: Optional[datetime] = None,
    include_graph: bool = True,
) -> str:
    """
    生成舆情分析报告
    :return: 生成的文件路径
    """
    doc = Document()

    # 封面
    title_para = doc.add_paragraph()
    title_run = title_para.add_run(title)
    title_run.font.size = Pt(24)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    date_para = doc.add_paragraph()
    date_para.add_run(f"生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if platforms:
        platform_para = doc.add_paragraph()
        platform_para.add_run(f"监控平台: {', '.join(platforms)}")
        platform_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    # 一、数据概览
    _add_heading(doc, "一、数据概览", level=1)

    total = get_post_count(platforms=platforms, start_time=start_time, end_time=end_time)
    sentiment_dist = get_sentiment_distribution(platforms=platforms, start_time=start_time, end_time=end_time)
    platform_dist = get_platform_distribution(start_time=start_time, end_time=end_time)

    kpis = {
        "总帖子数": total,
        "正面帖子": sentiment_dist.get("positive", 0),
        "负面帖子": sentiment_dist.get("negative", 0),
        "中性帖子": sentiment_dist.get("neutral", 0),
    }
    for plat, cnt in platform_dist.items():
        kpis[f"{plat}平台"] = cnt

    _add_kpi_table(doc, kpis)

    # 二、情感趋势
    _add_heading(doc, "二、情感趋势分析", level=1)

    trend = get_trend_by_time(
        interval="day",
        platforms=platforms,
        start_time=start_time,
        end_time=end_time,
    )
    if trend:
        indices = calculate_sentiment_index(trend)
        if indices:
            p = doc.add_paragraph()
            p.add_run("情感指数变化: ").bold = True
            latest = indices[-1]
            earliest = indices[0]
            change = latest["sentiment_index"] - earliest["sentiment_index"]
            direction = "上升" if change > 0 else "下降"
            p.add_run(f"从 {earliest['time_bucket']} 的 {earliest['sentiment_index']:.1f} "
                      f"变化至 {latest['time_bucket']} 的 {latest['sentiment_index']:.1f}，"
                      f"整体呈{direction}趋势（变化幅度: {change:+.1f}）。")

    # 三、热门话题
    _add_heading(doc, "三、热门话题", level=1)
    posts = get_posts(platforms=platforms, start_time=start_time, end_time=end_time, limit=1000)
    if posts:
        from analysis.topic_clustering import TopicClusterer
        texts = [p.get("content", "") for p in posts]
        clusterer = TopicClusterer(n_clusters=min(5, len(texts) // 10))
        topics = clusterer.fit(texts)

        for tid, info in topics.items():
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(f"话题 {tid}").bold = True
            p.add_run(f": 包含 {info['size']} 条帖子，关键词: {', '.join(info['keywords'][:5])}")

    # 四、传播分析
    if include_graph:
        _add_heading(doc, "四、传播网络分析", level=1)
        edges = get_edges(start_time=start_time, end_time=end_time)
        G = build_graph_from_posts(posts, edges, directed=True)
        stats = get_graph_stats(G)

        graph_kpis = {
            "网络节点数": stats["nodes"],
            "传播关系数": stats["edges"],
            "网络密度": stats["density"],
            "连通分量": stats["components"],
        }
        _add_kpi_table(doc, graph_kpis)

        influencers = detect_influencers(G, top_n=5)
        if influencers:
            _add_heading(doc, "关键传播节点", level=2)
            for inf in influencers:
                p = doc.add_paragraph(style="List Bullet")
                p.add_run(f"{inf['author']} ({inf['platform']})").bold = True
                p.add_run(f" - PageRank: {inf['pagerank']}, 互动量: {inf['engagement']}")

    # 五、告警摘要
    _add_heading(doc, "五、告警摘要", level=1)
    alerts = get_alerts(limit=20)
    if alerts:
        for alert in alerts:
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(f"[{alert['severity'].upper()}] {alert['alert_type']}").bold = True
            p.add_run(f": {alert['message']} ({alert['triggered_at']})")
    else:
        doc.add_paragraph("本周期内未触发告警。")

    # 六、建议措施
    _add_heading(doc, "六、建议措施", level=1)
    suggestions = []
    neg_ratio = sentiment_dist.get("negative", 0) / total * 100 if total > 0 else 0
    if neg_ratio > 30:
        suggestions.append("负面情感占比较高，建议启动危机公关预案，及时回应用户关切。")
    if platform_dist.get("微博", 0) > total * 0.5:
        suggestions.append("微博平台声量集中，建议优先在该平台发布正面引导内容。")
    if not suggestions:
        suggestions.append("当前舆情整体平稳，建议继续保持监测，定期发布正面内容。")

    for s in suggestions:
        doc.add_paragraph(s, style="List Bullet")

    # 保存
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"舆情报告_{timestamp}.docx"
    filepath = os.path.join(REPORTS_DIR, filename)
    os.makedirs(REPORTS_DIR, exist_ok=True)
    doc.save(filepath)
    return filepath


def get_report_list() -> List[Dict[str, str]]:
    """获取已生成的报告列表"""
    if not os.path.exists(REPORTS_DIR):
        return []

    reports = []
    for fname in sorted(os.listdir(REPORTS_DIR), reverse=True):
        if fname.endswith(".docx"):
            fpath = os.path.join(REPORTS_DIR, fname)
            reports.append({
                "filename": fname,
                "path": fpath,
                "created": datetime.fromtimestamp(os.path.getctime(fpath)).strftime("%Y-%m-%d %H:%M:%S"),
            })
    return reports
